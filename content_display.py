# Full Python code with Perceptual Hashing and modified output logic for all file types

# Section 1: Import Statements
import os
import json
import hashlib
import re
import math
import base64
import logging
import sys
import subprocess
import difflib
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
from collections import defaultdict, Counter
from importlib import import_module
# Import Pillow and ImageHash
try:
    from PIL import Image
    import imagehash
except ImportError:
    print("Please run 'pip install Pillow ImageHash' to use image similarity features.")
    sys.exit(1)

# Section 2: System Configuration
try:
    sys.stdout.reconfigure(encoding='utf-8')
except TypeError:
    pass

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Section 3: Dynamic Package Installation (Optional Helper)
def install_and_import(package_name: str, import_name: str = None):
    """Install and import a package if it's not available."""
    if import_name is None:
        import_name = package_name
    try:
        return import_module(import_name)
    except ImportError:
        logger.info(f"Installing {package_name}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
        return import_module(import_name)

# Section 4: MemorizationResult Class
class MemorizationResult:
    """A data class to store the results of a memorization check."""
    def __init__(self,
                 content_hash: str,
                 similarity_score: float,
                 is_memorized: bool,
                 confidence: float,
                 matched_content: str = None,
                 source_file: str = None,
                 detection_method: str = "unknown",
                 content_type: str = "text",
                 intersecting_content: str = None):
        self.content_hash = content_hash
        self.similarity_score = similarity_score
        self.is_memorized = is_memorized
        self.confidence = confidence
        self.matched_content = matched_content
        self.source_file = source_file
        self.detection_method = detection_method
        self.content_type = content_type
        self.intersecting_content = intersecting_content

# Section 5-12: EnhancedContentExtractor Class
class EnhancedContentExtractor:
    """Extracts and processes content from various file types."""
    def __init__(self):
        self.text_extensions = {'.txt', '.md', '.py', '.html', '.xml', '.json', '.csv'}
        self.image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}
        self.document_extensions = {'.pdf', '.docx', '.doc'}
        self.supported_extensions = self.text_extensions | self.image_extensions | self.document_extensions
        self.libraries_available = self._check_libraries()

    def _check_libraries(self) -> Dict[str, bool]:
        """Checks which optional libraries for file processing are installed."""
        available = {'PIL': True}
        try:
            import fitz
            available['fitz'] = True
        except ImportError:
            available['fitz'] = False
        try:
            from docx import Document
            available['docx'] = True
        except ImportError:
            available['docx'] = False
        return available

    def _extract_image_content(self, file_path: str) -> Dict[str, Any]:
        """Processes images using perceptual hashing (dHash)."""
        try:
            with Image.open(file_path) as img:
                perceptual_hash = imagehash.dhash(img)
                metadata = {
                    "file_size": Path(file_path).stat().st_size,
                    "extension": Path(file_path).suffix.lower(),
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode
                }
                return {
                    "content": str(perceptual_hash),
                    "content_type": "image",
                    "metadata": metadata
                }
        except Exception as e:
            logger.error(f"Error processing image {file_path} with PIL/ImageHash: {e}")
            return {"content": "", "content_type": "error", "metadata": {"error": str(e)}}
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File {file_path} not found")
        
        extension = path.suffix.lower()
        
        try:
            if extension in self.text_extensions:
                return self._extract_text_content(str(path), extension)
            elif extension in self.image_extensions:
                return self._extract_image_content(str(path))
            elif extension in self.document_extensions:
                return self._extract_document_content(str(path), extension)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return {"content": "", "content_type": "unsupported", "metadata": {}}
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return {"content": "", "content_type": "error", "metadata": {"error": str(e)}}

    def _extract_text_content(self, file_path: str, extension: str) -> Dict[str, Any]:
        if extension == '.json': content = self._read_json(file_path)
        elif extension == '.csv': content = self._read_csv(file_path)
        else: content = self._read_text_file(file_path)
        return {"content": content, "content_type": "text", "metadata": {"file_size": Path(file_path).stat().st_size, "extension": extension, "char_count": len(content)}}

    def _extract_document_content(self, file_path: str, extension: str) -> Dict[str, Any]:
        content = ""
        metadata = {"extension": extension, "file_size": Path(file_path).stat().st_size}
        if extension == '.pdf' and self.libraries_available.get('fitz', False):
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    content = "".join(page.get_text() for page in doc)
                    metadata["pages"] = len(doc)
            except Exception as e: logger.error(f"Error reading PDF {file_path}: {e}"); metadata["processing_error"] = str(e)
        elif extension in ['.docx', '.doc'] and self.libraries_available.get('docx', False):
            try:
                from docx import Document
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
                metadata["paragraphs"] = len(doc.paragraphs)
            except Exception as e: logger.error(f"Error reading DOCX {file_path}: {e}"); metadata["processing_error"] = str(e)
        if not content: content = self._read_text_file(file_path)
        return {"content": content, "content_type": "document", "metadata": metadata}

    def _read_text_file(self, file_path: str) -> str:
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f: return f.read()
            except (UnicodeDecodeError, UnicodeError): continue
        try:
            with open(file_path, 'rb') as f: return f.read().decode('utf-8', errors='ignore')
        except Exception as e: logger.error(f"Failed to read {file_path} with any encoding: {e}"); return ""

    def _read_json(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f: return json.dumps(json.load(f), indent=2)
        except Exception as e: logger.error(f"Error reading JSON {file_path}: {e}"); return ""

    def _read_csv(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f: return "".join(f.readlines()[:100])
        except Exception as e: logger.error(f"Error reading CSV {file_path}: {e}"); return ""

# Section 13-20: MultiModalMemorizationDetector Class
class MultiModalMemorizationDetector:
    """The core detection engine that compares input content against a training set."""
    def __init__(self, threshold: float = 0.85, image_threshold: int = 10):
        self.threshold = threshold
        self.image_threshold = image_threshold 
        self.training_data = {'text': [], 'image': [], 'document': []}
        self.content_hashes = {'text': set(), 'image': set(), 'document': set()}

    def add_training_content(self, content_data: Dict[str, Any], source: str = None):
        content = content_data.get('content')
        content_type = content_data.get('content_type', 'text')
        if not content or len(str(content).strip()) < 5: return
        content_hash_key = str(content) if content_type == 'image' else hashlib.sha256(str(content).encode('utf-8')).hexdigest()
        if content_hash_key in self.content_hashes.get(content_type, set()): return
        training_item = {'content': content, 'hash': content_hash_key, 'source': source, 'metadata': content_data.get('metadata', {})}
        if content_type == 'text': training_item['cleaned_text'] = self._clean_text(str(content))
        self.training_data[content_type].append(training_item)
        self.content_hashes[content_type].add(content_hash_key)
        
    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text); text = re.sub(r'[^\w\s.,;:!?-]', '', text); return text.strip().lower()

    def detect_memorization(self, content_data: Dict[str, Any]) -> MemorizationResult:
        content = content_data.get('content')
        content_type = content_data.get('content_type', 'text')
        if not content: return MemorizationResult("", 0.0, False, 0.0, "empty_content", content_type=content_type)
        
        if content_type == 'text':
            return self._detect_text_similarity(str(content), hashlib.sha256(str(content).encode('utf-8')).hexdigest())
        elif content_type == 'image':
            return self._detect_image_similarity(content_data)
        else: # Documents
            return self._detect_generic_similarity(str(content), hashlib.sha256(str(content).encode('utf-8')).hexdigest(), content_type)

    def _detect_text_similarity(self, content: str, content_hash: str) -> MemorizationResult:
        if content_hash in self.content_hashes.get('text', set()):
            item = next((i for i in self.training_data['text'] if i['hash'] == content_hash), None)
            return MemorizationResult(content_hash, 1.0, True, 1.0, str(item['content']), item['source'], "exact_hash_match", "text", str(item['content']))
        
        if not self.training_data['text']: return MemorizationResult(content_hash, 0.0, False, 1.0, "no_training_data", "text")
        cleaned_content = self._clean_text(content)
        content_words = set(cleaned_content.split())
        best_similarity, best_match = 0.0, None
        for item in self.training_data['text']:
            training_words = set(item.get('cleaned_text', '').split())
            if not content_words or not training_words: continue
            intersection = len(content_words.intersection(training_words))
            union = len(content_words.union(training_words))
            similarity = intersection / union if union > 0 else 0
            if similarity > best_similarity: best_similarity, best_match = similarity, item
        is_memorized = best_similarity >= self.threshold
        intersecting_content_str = None
        if best_match:
            matcher = difflib.SequenceMatcher(None, content, str(best_match['content']))
            matched_parts = [content[b.a:b.a + b.size].strip() for b in matcher.get_matching_blocks() if b.size > 0]
            if matched_parts: intersecting_content_str = " ... ".join(matched_parts)
        return MemorizationResult(content_hash, best_similarity, is_memorized, best_similarity if is_memorized else (1.0 - best_similarity), str(best_match['content']) if best_match else None, best_match['source'] if best_match else None, "word_jaccard_similarity", "text", intersecting_content_str)

    def _detect_image_similarity(self, content_data: Dict[str, Any]) -> MemorizationResult:
        """Compares perceptual hashes using Hamming distance."""
        input_hash_str = content_data.get('content')
        if not input_hash_str:
            return MemorizationResult("", 0.0, False, 1.0, "no_image_content", "image")

        input_hash = imagehash.hex_to_hash(input_hash_str)
        best_match_item = None
        smallest_distance = float('inf')

        for item in self.training_data['image']:
            training_hash = imagehash.hex_to_hash(item['content'])
            distance = input_hash - training_hash
            
            if distance < smallest_distance:
                smallest_distance = distance
                best_match_item = item

        is_memorized = smallest_distance <= self.image_threshold
        similarity_score = max(0, 1.0 - (smallest_distance / 32.0))

        return MemorizationResult(
            content_hash=input_hash_str,
            similarity_score=similarity_score,
            is_memorized=is_memorized,
            confidence=similarity_score,
            matched_content=f"Closest image match with Hamming distance: {smallest_distance}",
            source_file=best_match_item['source'] if best_match_item else None,
            detection_method="perceptual_hash_dhash",
            content_type="image"
        )
    
    def _detect_generic_similarity(self, content: str, content_hash: str, content_type: str) -> MemorizationResult:
        if content_hash in self.content_hashes.get(content_type, set()):
              item = next((i for i in self.training_data[content_type] if i['hash'] == content_hash), None)
              return MemorizationResult(content_hash, 1.0, True, 1.0, str(item['content']), item['source'], "exact_hash_match", content_type, str(item['content']))
        if not self.training_data.get(content_type): return MemorizationResult(content_hash, 0.0, False, 1.0, "no_training_data", content_type)
        best_similarity, best_match = 0.0, None
        if len(content) < 5: return MemorizationResult(content_hash, 0.0, False, 1.0, "content_too_short", content_type)
        content_ngrams = set(content[i:i+5] for i in range(len(content) - 4))
        for item in self.training_data[content_type]:
            training_content = str(item['content'])
            if len(training_content) < 5: continue
            training_ngrams = set(training_content[i:i+5] for i in range(len(training_content) - 4))
            intersection = len(content_ngrams.intersection(training_ngrams))
            union = len(content_ngrams.union(training_ngrams))
            similarity = intersection / union if union > 0 else 0
            if similarity > best_similarity: best_similarity, best_match = similarity, item
        is_memorized = best_similarity >= self.threshold
        
        # ***** START: ADDED LOGIC *****
        # This block was added to find and store the actual intersecting content for documents.
        intersecting_content_str = None
        if best_match:
            matcher = difflib.SequenceMatcher(None, content, str(best_match['content']))
            matched_parts = [content[b.a:b.a + b.size].strip() for b in matcher.get_matching_blocks() if b.size > 0]
            if matched_parts:
                intersecting_content_str = " ... ".join(matched_parts)
        # ***** END: ADDED LOGIC *****
        
        return MemorizationResult(content_hash, best_similarity, is_memorized, best_similarity if is_memorized else (1.0 - best_similarity), str(best_match['content']) if best_match else None, best_match['source'] if best_match else None, "ngram_jaccard_similarity", content_type, intersecting_content_str)

# Section 21-24: EnhancedLLMMemorizationSystem Class (Facade)
class EnhancedLLMMemorizationSystem:
    def __init__(self, threshold: float = 0.85, image_threshold: int = 10):
        self.content_extractor = EnhancedContentExtractor()
        self.detector = MultiModalMemorizationDetector(threshold, image_threshold)
        self.system_ready = False
        self._print_capabilities()
    def _print_capabilities(self): print("="*40 + "\n=== Memorization Detection System initialized ===\n" + "="*40)
    def load_training_data(self, data_directory: str):
        data_path = Path(data_directory)
        if not data_path.is_dir(): logger.error(f"Training data directory not found: {data_directory}"); return
        logger.info(f"Loading training data from '{data_directory}'...")
        all_files = [f for f in data_path.rglob('*') if f.is_file()]
        for i, file_path in enumerate(all_files):
            try:
                content_data = self.content_extractor.extract_content(str(file_path))
                if content_data and content_data['content_type'] != 'unsupported' and content_data['content'] != '': self.detector.add_training_content(content_data, source=str(file_path))
            except Exception as e: logger.error(f"Failed to process file {file_path}: {e}")
            if (i + 1) % 20 == 0: logger.info(f"   ... processed {i+1}/{len(all_files)} files.")
        self.system_ready = True
        logger.info(f"--- Training complete. {len(all_files)} files processed. ---")
    def detect_memorization(self, input_data: str) -> MemorizationResult:
        if not self.system_ready: logger.warning("System not trained."); return MemorizationResult("", 0.0, False, 0.0, "system_not_ready")
        input_path = Path(input_data)
        content_data = self.content_extractor.extract_content(str(input_path)) if input_path.is_file() else {'content': input_data, 'content_type': 'text'}
        if not content_data or not content_data.get('content'): return MemorizationResult("", 0.0, False, 1.0, "extraction_failed")
        return self.detector.detect_memorization(content_data)

# Section 25: Sample Data Creation
def create_sample_data() -> str:
    data_dir = Path("./sample_training_data"); data_dir.mkdir(exist_ok=True)
    texts = {"text_1.txt": "The quick brown fox jumps over the lazy dog.", "text_2.txt": "Python is a powerful, high-level, general-purpose programming language.", "text_3.txt": "Machine learning is a field of inquiry devoted to understanding and building methods that 'learn'."}
    for filename, content in texts.items():
        with open(data_dir / filename, 'w', encoding='utf-8') as f: f.write(content)
    with open(data_dir / "data.json", 'w', encoding='utf-8') as f: json.dump({"project": "Memorization Detection", "version": 1.0}, f, indent=2)
    with open(data_dir / "data.csv", 'w', encoding='utf-8') as f: f.write("id,name,role\n1,Alice,Engineer\n2,Bob,Analyst\n")
    print(f"Sample data created in '{data_dir.resolve()}'")
    return str(data_dir)

# Section 26-27: Main Demonstration and Script Entry Point
def main():
    # Set separate thresholds for text (0.75) and images (Hamming distance of 10)
    detector_system = EnhancedLLMMemorizationSystem(threshold=0.75, image_threshold=10)
    training_dir = Path("./training_data")
    if not training_dir.is_dir():
        print(f"'{training_dir}' not found.")
        training_dir = Path(create_sample_data())
    detector_system.load_training_data(str(training_dir))
    
    print("\n" + "="*20 + " Interactive Mode " + "="*22)
    print("Enter text or a file path (including images) to check. Type 'quit' to exit.")
    try:
        while True:
            user_input = input("\n> ").strip()
            if user_input.lower() in ['quit', 'exit', 'q']: break
            if not user_input: continue
            
            result = detector_system.detect_memorization(user_input)
            print(f"  -> Result: {'MEMORIZED' if result.is_memorized else 'Not Memorized'} "
                  f"(Score: {result.similarity_score:.2%}, Method: {result.detection_method})")

            # For text/documents, show closest match if similarity is > 1%
            if result.content_type in ['text', 'document']:
                if result.similarity_score > 0.01 and result.intersecting_content:
                    print(f"  -> Closest match found in source file: {result.source_file}")
                    print("-" * 15 + " Common Content Found " + "-" * 14)
                    print(result.intersecting_content)
                    print("-" * (32 + len(" Common Content Found ")))
            
            # For images, show closest match if similarity is > 1%
            elif result.content_type == 'image':
                if result.similarity_score > 0.01 and result.source_file:
                    print(f"  -> Closest match found in source file: {result.source_file}")
                    print(f"  -> Details: {result.matched_content}")

    except KeyboardInterrupt:
        print("\nExiting.")
    finally:
        print("Goodbye!")


if __name__ == "__main__":
    main()