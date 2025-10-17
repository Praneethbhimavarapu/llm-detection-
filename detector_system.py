import os
import json
import hashlib
import re
import logging
import sys
import difflib
from pathlib import Path
from typing import Dict, Any
import time # Import time for checking file modification

# --- Dependency Checks ---
try:
    from PIL import Image
    import imagehash
except ImportError:
    print("Please run 'pip install Pillow ImageHash' to use image similarity features.")
    sys.exit(1)

try:
    import librosa
    import numpy as np
except ImportError:
    print("Please run 'pip install librosa numpy' to use the new audio similarity features.")
    sys.exit(1)


logger = logging.getLogger("detector_system")
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s'
)


# [NO CHANGES TO THE FOLLOWING CLASSES: MemorizationResult, EnhancedContentExtractor, MultiModalMemorizationDetector]
# ... (The code for the first three classes remains identical to the previous version) ...
# For brevity, I will omit them here, but you should include the full, unchanged code for them.

class MemorizationResult:
    """A class to hold the results of a memorization detection check."""
    matched_content_path=""
    def __init__(self,
                 content_hash: str,
                 similarity_score: float,
                 is_memorized: bool,
                 confidence: float,
                 matched_content: str = None,
                 source_file: str = None,
                 detection_method: str = "unknown",
                 content_type: str = "text",
                 intersecting_content: str = None):
        self.content_hash = content_hash
        self.similarity_score = float(similarity_score)
        self.is_memorized = bool(is_memorized)
        self.confidence = float(confidence)
        self.matched_content = matched_content if matched_content is None else str(matched_content)
        self.source_file = source_file if source_file is None else str(source_file)
        self.detection_method = str(detection_method)
        self.content_type = str(content_type)
        self.intersecting_content = intersecting_content if intersecting_content is None else str(intersecting_content)

    def to_dict(self):
        """Converts the result object to a dictionary."""
        return {
            "content_hash": self.content_hash,
            "similarity_score": self.similarity_score,
            "is_memorized": self.is_memorized,
            "confidence": self.confidence,
            "matched_content": self.matched_content,
            "source_file": self.source_file,
            "detection_method": self.detection_method,
            "content_type": self.content_type,
            "intersecting_content": self.intersecting_content,
        }


class EnhancedContentExtractor:
    """Extracts content from various file types for similarity analysis."""
    def __init__(self):
        self.text_extensions = {'.txt', '.md', '.py', '.html', '.xml', '.json', '.csv'}
        self.image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}
        self.document_extensions = {'.pdf', '.docx', '.doc'}
        self.audio_extensions = {'.wav', '.mp3', '.flac', '.ogg', '.m4a'}
        self.libraries_available = self._check_libraries()

    def _check_libraries(self) -> Dict[str, bool]:
        """Checks for optional dependencies needed for certain file types."""
        available = {'PIL': True}
        try:
            import fitz
            available['fitz'] = True
        except ImportError:
            available['fitz'] = False
        try:
            from docx import Document
            available['docx'] = True
        except ImportError:
            available['docx'] = False
        try:
            import librosa
            available['librosa'] = True
        except ImportError:
            available['librosa'] = False
        return available

    def _extract_image_content(self, file_path: str) -> Dict[str, Any]:
        """Extracts a perceptual hash from an image file."""
        try:
            with Image.open(file_path) as img:
                perceptual_hash = imagehash.dhash(img)
                metadata = {
                    "file_size": Path(file_path).stat().st_size,
                    "extension": Path(file_path).suffix.lower(),
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode
                }
                return {
                    "content": perceptual_hash.__str__(),
                    "content_type": "image",
                    "metadata": metadata
                }
        except Exception as e:
            logger.error(f"Error processing image {file_path} with PIL/ImageHash: {e}")
            return {"content": "", "content_type": "error", "metadata": {"error": str(e)}}

    def _extract_audio_content(self, file_path: str) -> Dict[str, Any]:
        """Extracts an MFCC fingerprint from an audio file."""
        if not self.libraries_available.get('librosa'):
            logger.warning("Librosa not installed. Skipping audio content extraction.")
            return {"content": "", "content_type": "unsupported", "metadata": {}}
        try:
            y, sr = librosa.load(file_path, duration=120)
            mfccs = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
            mfcc_fingerprint = np.mean(mfccs, axis=1).tolist()
            
            metadata = {
                "file_size": Path(file_path).stat().st_size,
                "extension": Path(file_path).suffix.lower(),
                "duration_seconds": librosa.get_duration(y=y, sr=sr),
                "sample_rate": sr
            }
            return {
                "content": json.dumps(mfcc_fingerprint),
                "content_type": "audio",
                "metadata": metadata
            }
        except Exception as e:
            logger.error(f"Error processing audio {file_path} with Librosa: {e}")
            return {"content": "", "content_type": "error", "metadata": {"error": str(e)}}

    def extract_content(self, file_path_or_fileobj) -> Dict[str, Any]:
        """Main dispatcher to extract content based on file type."""
        path = None
        if hasattr(file_path_or_fileobj, "read") and hasattr(file_path_or_fileobj, "filename"):
            upload_dir = Path("./temp_uploads")
            upload_dir.mkdir(exist_ok=True)
            path = upload_dir / Path(file_path_or_fileobj.filename).name
            
            file_path_or_fileobj.seek(0)
            with open(path, "wb") as f:
                f.write(file_path_or_fileobj.read())
        else:
            path = Path(file_path_or_fileobj)

        if not path or not path.exists():
            raise FileNotFoundError(f"File {str(path)} not found")
        
        extension = path.suffix.lower()
        try:
            if extension in self.text_extensions:
                return self._extract_text_content(str(path), extension)
            elif extension in self.image_extensions:
                return self._extract_image_content(str(path))
            elif extension in self.document_extensions:
                return self._extract_document_content(str(path), extension)
            elif extension in self.audio_extensions:
                return self._extract_audio_content(str(path))
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return {"content": "", "content_type": "unsupported", "metadata": {}}
        except Exception as e:
            logger.error(f"Error extracting content from {str(path)}: {e}")
            return {"content": "", "content_type": "error", "metadata": {"error": str(e)}}

    def _extract_text_content(self, file_path: str, extension: str) -> Dict[str, Any]:
        """Extracts content from various text-based files."""
        if extension == '.json':
            content = self._read_json(file_path)
        elif extension == '.csv':
            content = self._read_csv(file_path)
        else:
            content = self._read_text_file(file_path)
        return {
            "content": content,
            "content_type": "text",
            "metadata": {
                "file_size": Path(file_path).stat().st_size,
                "extension": extension,
                "char_count": len(content)
            }
        }

    def _extract_document_content(self, file_path: str, extension: str) -> Dict[str, Any]:
        """Extracts text content from document files like PDF and DOCX."""
        content = ""
        metadata = {"extension": extension, "file_size": Path(file_path).stat().st_size}
        if extension == '.pdf' and self.libraries_available.get('fitz', False):
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    content = "".join(page.get_text() for page in doc)
                    metadata["pages"] = len(doc)
            except Exception as e:
                logger.error(f"Error reading PDF {file_path}: {e}")
                metadata["processing_error"] = str(e)
        elif extension in ['.docx', '.doc'] and self.libraries_available.get('docx', False):
            try:
                from docx import Document
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
                metadata["paragraphs"] = len(doc.paragraphs)
            except Exception as e:
                logger.error(f"Error reading DOCX {file_path}: {e}")
                metadata["processing_error"] = str(e)
        if not content:
            content = self._read_text_file(file_path)
        return {
            "content": content,
            "content_type": "document",
            "metadata": metadata
        }

    def _read_text_file(self, file_path: str) -> str:
        """Reads a text file with robust encoding handling."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        try:
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Failed to read {file_path} with any encoding: {e}")
            return ""

    def _read_json(self, file_path: str) -> str:
        """Reads and pretty-prints JSON content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON {file_path}: {e}")
            return ""

    def _read_csv(self, file_path: str) -> str:
        """Reads a sample of a CSV file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return "".join(f.readlines()[:100])
        except Exception as e:
            logger.error(f"Error reading CSV {file_path}: {e}")
            return ""


class MultiModalMemorizationDetector:
    """Detects memorization by comparing input content to a training dataset."""
    def __init__(self, threshold: float = 0.85, image_threshold: int = 10, audio_threshold: float = 25.0):
        self.threshold = float(threshold)
        self.image_threshold = int(image_threshold)
        self.audio_threshold = float(audio_threshold)
        self.training_data = {'text': [], 'image': [], 'document': [], 'audio': []}
        self.content_hashes = {'text': set(), 'image': set(), 'document': set(), 'audio': set()}

    def add_training_content(self, content_data: Dict[str, Any], source: str = None):
        """Adds processed content to the training database."""
        content = content_data.get('content', '')
        content_type = content_data.get('content_type', 'text')
        if not content or len(str(content).strip()) < 5:
            return

        cleaned_text = None
        if content_type in ['text', 'document']:
            cleaned_text = self._clean_text(str(content))
        
        if content_type in ['image', 'audio']:
            content_hash_key = str(content)
        elif cleaned_text:
            content_hash_key = hashlib.sha256(cleaned_text.encode('utf-8')).hexdigest()
        else:
            return

        if content_hash_key in self.content_hashes.get(content_type, set()):
            return
        
        if content_type in ['text', 'document']:
             if content_hash_key in self.content_hashes['text'] or content_hash_key in self.content_hashes['document']:
                 return

        training_item = {
            'content': content,
            'hash': content_hash_key,
            'source': source,
            'metadata': content_data.get('metadata', {})
        }
        if content_type in ['text', 'document']:
            training_item['cleaned_text'] = cleaned_text
            
        self.training_data[content_type].append(training_item)
        self.content_hashes[content_type].add(content_hash_key)

    def _clean_text(self, text: str) -> str:
        """Normalizes text for comparison."""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,;:!?-]', '', text)
        return text.strip().lower()

    def detect_memorization(self, content_data: Dict[str, Any]) -> MemorizationResult:
        """Main dispatcher for detecting memorization based on content type."""
        content = content_data.get('content', '')
        content_type = content_data.get('content_type', 'text')

        if not content or len(str(content).strip()) < 5:
            return MemorizationResult("", 0.0, False, 0.0, "empty_content", content_type=content_type)

        if content_type in ['text', 'document']:
            cleaned_content = self._clean_text(str(content))
            content_hash = hashlib.sha256(cleaned_content.encode('utf-8')).hexdigest()
            return self._detect_text_similarity(str(content), cleaned_content, content_hash, content_type)
        elif content_type == 'image':
            return self._detect_image_similarity(content_data)
        elif content_type == 'audio':
            return self._detect_audio_similarity(content_data)
        else:
            cleaned_content = self._clean_text(str(content))
            content_hash = hashlib.sha256(cleaned_content.encode('utf-8')).hexdigest()
            return self._detect_text_similarity(str(content), cleaned_content, content_hash, content_type)

    def _detect_text_similarity(self, original_content: str, cleaned_content: str, content_hash: str, content_type: str = "text") -> MemorizationResult:
        """Detects similarity for text-based content."""
        for text_type in ['text', 'document']:
            if content_hash in self.content_hashes.get(text_type, set()):
                item = next((i for i in self.training_data[text_type] if i['hash'] == content_hash), None)
                if item:
                    return MemorizationResult(content_hash, 1.0, True, 1.0, str(item['content']), item['source'], "exact_hash_match", content_type, str(item['content']))

        searchable_text_data = self.training_data['text'] + self.training_data['document']
        if not searchable_text_data:
            return MemorizationResult(content_hash, 0.0, False, 1.0, "no_training_data", content_type)

        content_words = set(cleaned_content.split())
        best_similarity, best_match = 0.0, None
        
        for item in searchable_text_data:
            training_words = set(item.get('cleaned_text', '').split())
            if not content_words or not training_words:
                continue
            
            intersection = len(content_words.intersection(training_words))
            union = len(content_words.union(training_words))
            similarity = intersection / union if union > 0 else 0
            
            if similarity > best_similarity:
                best_similarity, best_match = similarity, item

        is_memorized = best_similarity >= self.threshold
        formatted_matched_content = None
        
        if best_match:
            matcher = difflib.SequenceMatcher(None, original_content.lower(), best_match['content'].lower())
            
            matched_parts = []
            for block in matcher.get_matching_blocks():
                part = best_match['content'][block.b : block.b + block.size]
                if part.strip():
                    matched_parts.append(part)
            
            if matched_parts:
                formatted_matched_content = " [...] ".join(matched_parts)

        return MemorizationResult(
            content_hash,
            best_similarity,
            is_memorized,
            best_similarity if is_memorized else (1.0 - best_similarity),
            formatted_matched_content,
            best_match['source'] if best_match else None,
            "word_jaccard_similarity + difflib",
            content_type,
            formatted_matched_content
        )

    def _detect_image_similarity(self, content_data: Dict[str, Any]) -> MemorizationResult:
        """Detects similarity for image content using perceptual hashes."""
        input_hash_str = content_data.get('content')
        if not input_hash_str:
            return MemorizationResult("", 0.0, False, 1.0, "no_image_content", "image")

        try:
            input_hash = imagehash.hex_to_hash(input_hash_str)
            best_match_item = None
            smallest_distance = float('inf')
            for item in self.training_data['image']:
                training_hash = imagehash.hex_to_hash(item['content'])
                distance = input_hash - training_hash
                if distance < smallest_distance:
                    smallest_distance = distance
                    best_match_item = item
            is_memorized = smallest_distance <= self.image_threshold
            similarity_score = max(0.0, 1.0 - (float(smallest_distance) / 32.0))
            matched_content_path = best_match_item['source'] if best_match_item else None

            return MemorizationResult(
                content_hash=input_hash_str,
                similarity_score=similarity_score,
                is_memorized=is_memorized,
                confidence=similarity_score,
                matched_content=matched_content_path,
                source_file=matched_content_path,
                detection_method="perceptual_hash_dhash",
                content_type="image"
            )
        except Exception as ex:
            logger.error(f"Error during image similarity computation: {ex}")
            return MemorizationResult("", 0.0, False, 0.0, "image_similarity_failed", "image")

    def _detect_audio_similarity(self, content_data: Dict[str, Any]) -> MemorizationResult:
        """Detects similarity for audio content using MFCC fingerprints."""
        input_fingerprint_str = content_data.get('content')
        if not input_fingerprint_str:
            return MemorizationResult("", 0.0, False, 1.0, "no_audio_content", "audio")

        try:
            input_fingerprint = np.array(json.loads(input_fingerprint_str))
            best_match_item = None
            smallest_distance = float('inf')
            
            for item in self.training_data['audio']:
                training_fingerprint = np.array(json.loads(item['content']))
                distance = np.linalg.norm(input_fingerprint - training_fingerprint)
                if distance < smallest_distance:
                    smallest_distance = distance
                    best_match_item = item

            is_memorized = smallest_distance <= self.audio_threshold
            normalization_factor = self.audio_threshold * 2 
            similarity_score = max(0.0, 1.0 - (smallest_distance / normalization_factor))
            
            matched_content_path = best_match_item['source'] if best_match_item else None

            return MemorizationResult(
                content_hash=input_fingerprint_str,
                similarity_score=similarity_score,
                is_memorized=is_memorized,
                confidence=similarity_score,
                matched_content=f"Euclidean Distance: {smallest_distance:.2f}",
                source_file=matched_content_path,
                detection_method="audio_mfcc_euclidean_distance",
                content_type="audio"
            )
        except Exception as ex:
            logger.error(f"Error during audio similarity computation: {ex}")
            return MemorizationResult("", 0.0, False, 0.0, "audio_similarity_failed", "audio")

            
class EnhancedLLMMemorizationSystem:
    """Top-level class to manage and run the memorization detection system."""
    def __init__(self, threshold: float = 0.85, image_threshold: int = 10, audio_threshold: float = 25.0, cache_path: str = "training_cache.json"):
        self.content_extractor = EnhancedContentExtractor()
        self.detector = MultiModalMemorizationDetector(threshold, image_threshold, audio_threshold)
        self.system_ready = False
        # [ADDED] Path to the cache file
        self.cache_path = Path(cache_path)

    # [ADDED] Method to load the cache from disk
    def _load_cache(self) -> Dict:
        if self.cache_path.exists():
            logger.info(f"Loading existing cache from '{self.cache_path}'")
            with open(self.cache_path, 'r') as f:
                try:
                    return json.load(f)
                except json.JSONDecodeError:
                    logger.warning("Cache file is corrupted. Starting fresh.")
                    return {}
        return {}

    # [ADDED] Method to save the cache to disk
    def _save_cache(self, cache_data: Dict):
        logger.info(f"Saving updated cache to '{self.cache_path}'")
        with open(self.cache_path, 'w') as f:
            json.dump(cache_data, f)

    # [MODIFIED] The load_training_data method is now much faster on subsequent runs
    def load_training_data(self, data_directory: str):
        """Loads and processes all files from a directory, using a cache for speed."""
        data_path = Path(data_directory)
        if not data_path.is_dir():
            logger.error(f"Training data directory not found: {data_directory}")
            return

        cache_data = self._load_cache()
        
        logger.info(f"Scanning and processing training data from '{data_directory}'...")
        all_files = [f for f in data_path.rglob('*') if f.is_file()]
        
        files_processed = 0
        files_from_cache = 0

        for file_path in all_files:
            file_str = str(file_path)
            try:
                # Create a unique key for the cache based on path and modification time
                mod_time = file_path.stat().st_mtime
                cache_key = f"{file_str}|{mod_time}"

                # Check if a valid result is already in the cache
                if cache_key in cache_data:
                    content_data = cache_data[cache_key]
                    files_from_cache += 1
                else:
                    # If not in cache, do the slow extraction and store the result
                    content_data = self.content_extractor.extract_content(file_str)
                    cache_data[cache_key] = content_data
                    files_processed += 1
                
                # Add the content (from cache or new) to the detector
                if content_data and content_data['content_type'] != 'unsupported' and content_data['content'] != '':
                    self.detector.add_training_content(content_data, source=file_str)

            except Exception as e:
                logger.error(f"Failed to process file {file_path}: {e}")

        if files_processed > 0: # Only save the cache if there were new files
             self._save_cache(cache_data)

        self.system_ready = True
        
        text_count = len(self.detector.training_data['text'])
        doc_count = len(self.detector.training_data['document'])
        img_count = len(self.detector.training_data['image'])
        audio_count = len(self.detector.training_data['audio'])
        
        logger.info("--- Training complete ---")
        logger.info(f"Loaded {files_from_cache} files from cache.")
        logger.info(f"Newly processed {files_processed} files.")
        logger.info(f"Total items in database: {text_count} text, {doc_count} docs, {img_count} images, {audio_count} audio files.")

    def detect_memorization(self, input_data: Any) -> MemorizationResult:
        """Detects memorization in a given input (file path, file object, or raw text)."""
        if not self.system_ready:
            logger.warning("System not trained.")
            return MemorizationResult("", 0.0, False, 0.0, "system_not_ready")
            
        content_data = None
        if hasattr(input_data, "read") and hasattr(input_data, "filename"):
            content_data = self.content_extractor.extract_content(input_data)
        elif isinstance(input_data, str):
            input_path = Path(input_data)
            if input_path.is_file():
                content_data = self.content_extractor.extract_content(str(input_path))
            else:
                content_data = {'content': input_data, 'content_type': 'text'}
        else:
             return MemorizationResult("", 0.0, False, 1.0, "unsupported_input_type")

        if not content_data or not content_data.get('content'):
            return MemorizationResult("", 0.0, False, 1.0, "extraction_failed")
            
        return self.detector.detect_memorization(content_data)
# ```

# **How to use the optimized code:**
# 1.  Run your training script as usual.
# 2.  The first run will be slow, just like before. But it will create a new file named `training_cache.json` in your directory.
# 3.  Run the *exact same script again*. You will see that it loads almost instantly, with a log message saying "Loaded X files from cache."

# ### Solution 2: Tune Audio Processing Parameters

# You can make the audio processing itself faster by telling `librosa` to do less work. This involves a trade-off: you gain speed but slightly reduce the accuracy or detail of the audio fingerprint.

# In the `EnhancedContentExtractor`, find the `_extract_audio_content` method and modify the `librosa.load` call.

# **Change this:**
# ```python
# y, sr = librosa.load(file_path, duration=120)
# ```

# **To this (a faster version):**
# ```python
# # Down-sample to 22.05kHz and process only the first 60 seconds
# y, sr = librosa.load(file_path, sr=22050, duration=60)

